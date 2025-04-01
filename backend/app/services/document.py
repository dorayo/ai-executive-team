import os
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import Document
from app.knowledge.vector_store import add_documents

# 支持的文档类型
SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "text/plain": "text",
    "text/html": "html",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx"
}

async def save_upload_file(upload_file: UploadFile) -> Tuple[str, str]:
    """
    保存上传文件到磁盘
    
    Args:
        upload_file: 上传的文件
        
    Returns:
        保存的文件路径和内容类型
    """
    content_type = upload_file.content_type
    
    # 检查文件类型
    if content_type not in SUPPORTED_CONTENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: {content_type}。支持的类型: {', '.join(SUPPORTED_CONTENT_TYPES.keys())}"
        )
    
    # 创建文件存储目录
    upload_dir = os.path.join(settings.UPLOAD_DIR, datetime.now().strftime("%Y%m"))
    os.makedirs(upload_dir, exist_ok=True)
    
    # 生成唯一文件名
    file_extension = os.path.splitext(upload_file.filename)[1]
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(upload_dir, unique_filename)
    
    # 保存文件
    with open(file_path, "wb") as f:
        content = await upload_file.read()
        f.write(content)
    
    return file_path, content_type

def extract_text_from_file(file_path: str, content_type: str) -> str:
    """
    从文件中提取文本
    
    Args:
        file_path: 文件路径
        content_type: 文件内容类型
        
    Returns:
        提取的文本内容
    """
    file_type = SUPPORTED_CONTENT_TYPES.get(content_type)
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "text":
        return extract_text_from_text(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "html":
        return extract_text_from_html(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {content_type}")

def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件中提取文本"""
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text = ""
        
        for page in reader.pages:
            text += page.extract_text() + "\n\n"
            
        return text
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"PDF文本提取失败: {str(e)}"
        )

def extract_text_from_text(file_path: str) -> str:
    """从文本文件中提取文本"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试不同的编码
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"文本文件读取失败: {str(e)}"
            )

def extract_text_from_docx(file_path: str) -> str:
    """从DOCX文件中提取文本"""
    try:
        import docx
        
        doc = docx.Document(file_path)
        text = ""
        
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        return text
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"DOCX文本提取失败: {str(e)}"
        )

def extract_text_from_html(file_path: str) -> str:
    """从HTML文件中提取文本"""
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            return soup.get_text()
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"HTML文本提取失败: {str(e)}"
        )

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """
    将文本分割成较小的块
    
    Args:
        text: 完整文本
        chunk_size: 每个块的最大大小(字符数)
        chunk_overlap: 块之间的重叠(字符数)
        
    Returns:
        文本块列表
    """
    if not text:
        return []
        
    chunks = []
    start = 0
    
    while start < len(text):
        # 找到块的结束位置
        end = min(start + chunk_size, len(text))
        
        # 如果不是最后一个块，尝试在句子边界处分割
        if end < len(text):
            # 向后寻找句子结束标记
            for i in range(end, max(end - chunk_overlap, start), -1):
                if i < len(text) and text[i] in ['.', '!', '?', '\n'] and (i + 1 >= len(text) or text[i + 1] == ' '):
                    end = i + 1
                    break
        
        # 添加块
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # 向前移动，考虑重叠
        start = end - chunk_overlap if end - chunk_overlap > start else start + 1
    
    return chunks

async def create_document(
    db: Session,
    file: UploadFile,
    title: str,
    description: Optional[str] = None,
    user_id: Optional[int] = None
) -> Document:
    """
    创建新文档，处理文件，并存储向量
    
    Args:
        db: 数据库会话
        file: 上传的文件
        title: 文档标题
        description: 文档描述
        user_id: 上传用户ID
        
    Returns:
        创建的文档对象
    """
    # 保存文件到磁盘
    file_path, content_type = await save_upload_file(file)
    
    # 创建文档记录
    doc = Document(
        title=title,
        description=description,
        file_path=file_path,
        content_type=content_type,
        uploaded_by=user_id,
        processing_status="pending"
    )
    
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    # 提取文本
    try:
        # 更新状态为处理中
        doc.processing_status = "processing"
        db.commit()
        
        text = extract_text_from_file(file_path, content_type)
        
        # 为较小的文档存储内容
        if len(text) <= 100000:  # 大约10万字符(约20页)
            doc.content = text
            db.commit()
        
        # 将文本分块
        chunks = chunk_text(text)
        
        if chunks and settings.PINECONE_API_KEY:
            try:
                # 为每个块准备元数据
                metadatas = [{
                    "document_title": title,
                    "document_id": doc.id,
                    "chunk_index": i
                } for i in range(len(chunks))]
                
                # 将文档添加到向量存储
                vector_ids = add_documents(chunks, metadatas, doc.id)
                
                # 存储向量ID
                doc.vector_ids = ",".join(vector_ids)
                doc.processing_status = "completed"
                db.commit()
            except Exception as e:
                # 捕获向量存储异常但继续
                doc.processing_status = "partial"
                doc.processing_error = f"向量存储失败，但文本内容已保存: {str(e)}"
                db.commit()
                print(f"向量存储失败: {str(e)}")
        else:
            # 没有块或没有API密钥，但文本处理成功
            doc.processing_status = "text_only"
            if not settings.PINECONE_API_KEY:
                doc.processing_error = "未配置Pinecone API密钥，只保存了文本内容"
            db.commit()
    except Exception as e:
        # 如果处理失败，仍保留文档记录，但设置错误标记
        doc.processing_status = "error"
        doc.processing_error = str(e)
        db.commit()
        
        raise HTTPException(
            status_code=500,
            detail=f"文档处理失败: {str(e)}"
        )
    
    return doc

def get_document(db: Session, document_id: int) -> Optional[Document]:
    """获取文档"""
    return db.query(Document).filter(Document.id == document_id).first()

def get_documents(db: Session, skip: int = 0, limit: int = 100) -> List[Document]:
    """获取文档列表"""
    return db.query(Document).order_by(Document.created_at.desc()).offset(skip).limit(limit).all()

def delete_document(db: Session, document_id: int) -> bool:
    """删除文档"""
    from app.knowledge.vector_store import delete_document_vectors
    
    doc = get_document(db, document_id)
    if not doc:
        return False
    
    # 删除向量存储中的文档向量
    if doc.vector_ids:
        delete_document_vectors(doc.id)
    
    # 删除文件
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except Exception:
            # 即使文件删除失败，也继续删除数据库记录
            pass
    
    # 删除数据库记录
    db.delete(doc)
    db.commit()
    
    return True 